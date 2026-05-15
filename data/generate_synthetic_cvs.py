import argparse
import random
import os
from pathlib import Path
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch

FIRST_NAMES = ["Jan", "Petr", "Pavel", "Tomáš", "Martin", "Jakub", "Lukáš", "David", "Michal", "Filip",
               "Anna", "Marie", "Eva", "Petra", "Jana", "Lucie", "Kateřina", "Lenka", "Martina", "Veronika"]
LAST_NAMES = ["Novák", "Svoboda", "Novotný", "Dvořák", "Černý", "Procházka", "Kučera", "Veselý", "Horák", "Němec",
              "Marek", "Pokorný", "Pospíšil", "Hájek", "Král", "Jelínek", "Růžička", "Beneš", "Fiala", "Sedláček"]

SKILLS = [
    "Python", "Java", "JavaScript", "React", "Angular", "Node.js", "AWS", "Azure", "GCP", "Docker",
    "Kubernetes", "SQL", "NoSQL", "PostgreSQL", "MongoDB", "Terraform", "CI/CD", "Machine Learning",
    "Data Engineering", "DevOps", "Cloud Architecture", "Microservices", "REST APIs", "GraphQL",
    "Agile", "Scrum", "Leadership", "Team Management", "Salesforce", "CRM", "Customer Service",
    "Content Creation", "Copywriting", "SEO", "SEM", "Financial Analysis", "Budgeting", "Excel",
    "Project Management", "UX Design", "Adobe Photoshop", "Figma", "Illustrator", "Quality Assurance",
    "Network Administration", "Helpdesk Support", "Healthcare Compliance", "Clinical Research",
    "Business Development", "Email Marketing", "Social Media", "Human Resources"
]

ROLE_POOLS = {
    "IT": [
        ("Junior Developer", 1, 2),
        ("Software Developer", 2, 4),
        ("Senior Developer", 4, 7),
        ("Lead Developer", 6, 10),
        ("Software Architect", 8, 12),
        ("Engineering Manager", 7, 15),
        ("Principal Engineer", 10, 15),
        ("IT Support Specialist", 1, 4),
    ],
    "Business": [
        ("Marketing Assistant", 1, 3),
        ("Sales Representative", 1, 4),
        ("Business Analyst", 2, 5),
        ("Project Coordinator", 1, 4),
        ("Product Manager", 3, 8),
        ("Operations Manager", 5, 12),
    ],
    "Design": [
        ("Graphic Designer", 1, 3),
        ("UX Designer", 2, 5),
        ("Visual Designer", 1, 4),
        ("Creative Director", 6, 12),
        ("Art Director", 5, 10),
    ],
    "Healthcare": [
        ("Healthcare Assistant", 1, 3),
        ("Clinical Researcher", 2, 5),
        ("Nursing Specialist", 3, 7),
        ("Healthcare Manager", 5, 12),
    ],
    "Finance": [
        ("Accounting Assistant", 1, 3),
        ("Financial Analyst", 2, 5),
        ("Tax Specialist", 2, 5),
        ("Senior Accountant", 4, 8),
        ("Finance Manager", 6, 12),
    ],
}

COMPANIES = [
    "TechCorp", "DataSystems", "CloudWorks", "InnoSoft", "DevHub", "CodeFactory", "ByteLabs",
    "AgileWorks", "SmartSolutions", "DigitalForge", "WebMasters", "AppBuilders", "RetailFront", "HealthWay",
    "MarketPulse", "FinancePro", "CreativeStudio", "GreenLogistics", "UrbanServices"
]

EDUCATION = [
    "High School Diploma",
    "Vocational Certificate in IT Support",
    "Bachelor of Science in Computer Science",
    "Bachelor of Arts in Marketing",
    "Bachelor of Science in Graphic Design",
    "Bachelor of Science in Nursing",
    "Master of Science in Software Engineering",
    "Master of Business Administration",
    "Master of Science in Data Analytics",
    "PhD in Computer Science"
]

SUMMARY_TEMPLATES = {
    "IT": [
        "Experienced professional with {years} years in software development and systems engineering.",
        "Technology specialist focused on delivering stable applications and cloud infrastructure.",
        "Developer with strong skills in building scalable solutions and automating operations.",
    ],
    "Business": [
        "Business professional with {years} years driving sales, operations and process improvements.",
        "Analytical thinker with experience in project delivery and stakeholder communication.",
        "Result-oriented contributor helping teams meet targets and improve customer outcomes.",
    ],
    "Design": [
        "Creative designer with {years} years producing user-centered visuals and digital experiences.",
        "Design specialist blending brand storytelling with practical execution.",
        "Visual communicator experienced in digital campaigns and product interfaces.",
    ],
    "Healthcare": [
        "Healthcare professional with {years} years supporting patient care and medical operations.",
        "Clinically minded specialist focused on process quality and compliance.",
        "Patient-oriented practitioner with experience in multidisciplinary teams.",
    ],
    "Finance": [
        "Finance professional with {years} years analyzing performance and supporting budgeting decisions.",
        "Detail-oriented analyst experienced in financial reporting and business planning.",
        "Accounting specialist supporting accurate financial operations and risk management.",
    ],
}


def _build_experience_path(path_roles, target_years):
    experience = []
    current_year = 2024
    years_exp = 0

    while years_exp < target_years and current_year > 2010:
        role_info = random.choice(path_roles)
        min_dur, max_dur = role_info[1], role_info[2]
        remaining = target_years - years_exp
        duration = random.randint(min_dur, min(max_dur, max(min_dur, remaining)))
        start_year = current_year - duration
        experience.append({
            "role": role_info[0],
            "company": random.choice(COMPANIES),
            "start": start_year,
            "end": current_year,
            "duration": duration,
        })
        years_exp += duration
        current_year = start_year
        if len(experience) >= 4:
            break

    return experience, years_exp


def generate_cv_data():
    name = f"{random.choice(FIRST_NAMES)} {random.choice(LAST_NAMES)}"
    industry = random.choice(list(ROLE_POOLS.keys()))
    path_roles = ROLE_POOLS[industry]

    career_type = random.choices(
        ["junior", "mid", "senior"],
        weights=[40, 40, 20],
        k=1,
    )[0]

    if career_type == "junior":
        target_years = random.randint(1, 3)
    elif career_type == "mid":
        target_years = random.randint(3, 6)
    else:
        target_years = random.randint(6, 15)

    experience, years_exp = _build_experience_path(path_roles, target_years)
    if not experience:
        experience, years_exp = _build_experience_path(path_roles, max(1, target_years))

    if career_type == "junior":
        num_skills = random.randint(3, 7)
    elif career_type == "mid":
        num_skills = random.randint(5, 10)
    else:
        num_skills = random.randint(8, 14)

    skills = random.sample(SKILLS, min(len(SKILLS), num_skills))
    education = random.choice(EDUCATION)

    summary = random.choice(SUMMARY_TEMPLATES[industry]).format(years=years_exp)
    if career_type == "junior":
        summary = f"Motivated {industry.lower()} professional with {years_exp} years of practical experience. {summary}"
    elif career_type == "senior":
        summary = f"Senior {industry.lower()} specialist with {years_exp} years of proven results. {summary}"

    return {
        "name": name,
        "summary": summary,
        "experience": experience,
        "skills": skills,
        "education": education,
    }


def generate_docx(cv_data, output_path):
    doc = Document()
    
    doc.add_heading(cv_data["name"], 0)
    
    doc.add_heading("Summary", level=2)
    doc.add_paragraph(cv_data["summary"])
    
    doc.add_heading("Experience", level=2)
    for exp in cv_data["experience"]:
        doc.add_paragraph(
            f"{exp['role']} at {exp['company']} ({exp['start']}–{exp['end']})",
            style='List Bullet'
        )
    
    doc.add_heading("Skills", level=2)
    doc.add_paragraph(", ".join(cv_data["skills"]))
    
    doc.add_heading("Education", level=2)
    doc.add_paragraph(cv_data["education"])
    
    doc.save(output_path)


def generate_pdf(cv_data, output_path):
    c = canvas.Canvas(output_path, pagesize=letter)
    width, height = letter
    
    y = height - inch
    
    c.setFont("Helvetica-Bold", 18)
    c.drawString(inch, y, cv_data["name"])
    y -= 0.5 * inch
    
    c.setFont("Helvetica-Bold", 14)
    c.drawString(inch, y, "Summary")
    y -= 0.3 * inch
    
    c.setFont("Helvetica", 11)
    c.drawString(inch, y, cv_data["summary"])
    y -= 0.5 * inch
    
    c.setFont("Helvetica-Bold", 14)
    c.drawString(inch, y, "Experience")
    y -= 0.3 * inch
    
    c.setFont("Helvetica", 11)
    for exp in cv_data["experience"]:
        text = f"• {exp['role']} at {exp['company']} ({exp['start']}–{exp['end']})"
        c.drawString(inch, y, text)
        y -= 0.25 * inch
        if y < inch:
            break
    
    y -= 0.2 * inch
    c.setFont("Helvetica-Bold", 14)
    c.drawString(inch, y, "Skills")
    y -= 0.3 * inch
    
    c.setFont("Helvetica", 11)
    skills_text = ", ".join(cv_data["skills"])
    c.drawString(inch, y, skills_text[:80])
    if len(skills_text) > 80:
        y -= 0.25 * inch
        c.drawString(inch, y, skills_text[80:])
    
    y -= 0.5 * inch
    c.setFont("Helvetica-Bold", 14)
    c.drawString(inch, y, "Education")
    y -= 0.3 * inch
    
    c.setFont("Helvetica", 11)
    c.drawString(inch, y, cv_data["education"])
    
    c.save()


def main():
    parser = argparse.ArgumentParser(description="Generate synthetic CVs for testing")
    parser.add_argument("-n", "--number", type=int, default=20, help="Number of CVs to generate (default: 20)")
    parser.add_argument("-o", "--output", type=str, default="sample", help="Output directory (default: data/sample)")
    args = parser.parse_args()
    
    output_dir = Path(args.output)
    output_dir.mkdir(parents=True, exist_ok=True)
    
    print(f"Generating {args.number} synthetic CVs in {output_dir}...")
    
    for i in range(args.number):
        cv_data = generate_cv_data()
        
        if i % 2 == 0:
            output_path = output_dir / f"cv_{i+1:03d}.docx"
            generate_docx(cv_data, str(output_path))
        else:
            output_path = output_dir / f"cv_{i+1:03d}.pdf"
            generate_pdf(cv_data, str(output_path))
        
        print(f"Created: {output_path}")
    
    print(f"\nSuccessfully generated {args.number} CVs!")


if __name__ == "__main__":
    main()
